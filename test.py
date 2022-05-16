from resume_parser import ResumeParser
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import codecs
import csv
resume_path = './resumes1'
information = []
resume_list = os.listdir(resume_path)
# resume_list = ['管理岗+陈凯欣+北京交通大学+新闻传播学.pdf']
for i in range(len(resume_list)):
    print('正在提取' + resume_list[i])
    resume_name = resume_list[i]
    print('=================================================')
    number = ''
    name_zh = ''
    dob_zh = ''
    gender_zh = ''
    political_face = ''
    highest_degree = ''
    undergraduate_school = ''
    undergraduate_major = ''
    master_school = ''
    master_major = ''
    doctoral_school = ''
    doctoral_major = ''
    is_beijing = ''
    hometown = ''
    phone_number = ''
    mail = ''

    data = ResumeParser('resumes1/' + resume_list[i]).get_extracted_data()
    # pdftext = ResumeParser('resumes/' + resume_list[i]).get_pdftext()
    # clean_text = ResumeParser('resumes/' + resume_list[i]).get_clean_pdftext()
    number = str(i)
    name_zh = data['name_zh']
    dob_zh = data['dob_zh']
    gender_zh = data['gender_zh']
    political_face = data['political_face']
    highest_degree = data['highest_degree']
    undergraduate_school = data['undergraduate_school']
    undergraduate_major = data['undergraduate_major']
    master_school = data['master_school']
    master_major = data['master_major']
    doctoral_school = data['doctoral_school']
    doctoral_major = data['doctoral_major']
    is_beijing = data['is_beijing']
    hometown = data['hometown']
    phone_number = data['phone_number']
    mail = data['mail']

    print('姓名' + str(name_zh))
    print('性别' + str(gender_zh))
    print('生日' + str(dob_zh))
    print('政治面貌' + str(political_face))
    print('学历' + str(highest_degree))
    print('本科' + str(undergraduate_school))
    print('本科' + str(undergraduate_major))
    print('硕士' + str(master_school))
    print('硕士' + str(master_major))
    print('博士' + str(doctoral_school))
    print('博士' + str(doctoral_major))
    print('京籍' + str(is_beijing))
    print('籍贯' + str(hometown))
    print('电话' + str(phone_number))
    print('邮箱' + str(mail))
    number = str(i)

    information.append([number, name_zh, gender_zh, dob_zh, political_face, highest_degree, undergraduate_school, undergraduate_major, master_school, master_major, doctoral_school, doctoral_major, is_beijing, hometown, phone_number, mail])

    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    Head = document.add_heading('', level=1)
    Head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = Head.add_run("个人简历")
    run.font.name=u'Cambria'
    run.font.color.rgb = RGBColor(0,0,0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    document.add_paragraph('姓名 : ' + str(name_zh))
    document.add_paragraph('性别 : ' + str(gender_zh))
    document.add_paragraph('出生年月日 : ' + str(dob_zh))
    document.add_paragraph('政治面貌 : ' + str(political_face))
    document.add_paragraph('学历（最高学历）: ' + str(highest_degree))
    document.add_paragraph('本科毕业院校 : ' + str(undergraduate_school))
    document.add_paragraph('本科所学专业 : ' + str(undergraduate_major))
    document.add_paragraph('硕士毕业院校 : ' + str(master_school))
    document.add_paragraph('硕士所学专业 : ' + str(master_major))
    document.add_paragraph('博士毕业院校 : ' + str(doctoral_school))
    document.add_paragraph('博士所学专业 : ' + str(doctoral_major))
    document.add_paragraph('是否京籍 : ' + str(is_beijing))
    document.add_paragraph('籍贯 : ' + str(hometown))
    document.add_paragraph('联系方式（手机号）: ' + str(phone_number))
    document.add_paragraph('电子邮箱 : ' + str(mail))
    save_path = './output_docx/' + resume_name + 'after_extra.docx'
    document.add_page_break()
    document.save(save_path)

header = ['编号', '姓名', '性别' ,'出生年月日', '政治面貌', '学历', '本科毕业院校', '本科专业', '硕士毕业院校', '硕士专业', '博士毕业院校', '博士专业', '是否京籍学生', '籍贯', '手机号', '邮箱']
output_file = './output.csv'
f = codecs.open(output_file, 'a+', 'gbk')
writer = csv.writer(f)
writer.writerow(header)
try:
    writer.writerows(information)
except UnicodeEncodeError as e:
    print(e)

f.close()
print('csv文件已生成，为在项目路径下的output.csv')
